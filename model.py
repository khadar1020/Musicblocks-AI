import torch
import os
import google, googlesearch
from googlesearch import search
from transformers import pipeline
from langchain.prompts import PromptTemplate
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.chains import RetrievalQA
import chainlit as cl
from langchain_community.llms import Ollama
from fuzzywuzzy import fuzz
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.output_parsers import StrOutputParser
from langchain_core.messages import AIMessage, HumanMessage
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DB_CHROMA_PATH = 'vectorstore/db_chroma'

custom_prompt_template = """
1. 8 to 10 years children should also be capable of understanding your answer so just divide your answer into points with numbers or stars or some decorations(Important - Don't write in answer this point) and explain the answer in detail so the children can feel it interactive🎵🎹.
2. (Important)If any question the user asks which is not related to music or music blocks then just say This is the music blocks or generating a lesson plan or lessons chatbot only answers music blocks questions polietly. 
3. When the user asks a question like generate a lesson plan of any particular topic then follow the structure provided. 
4. (Important)If the user asks any question related to maths caluculation also say I will have to say I will not answer this question this si music blocks chatbot please ank questions related to music and exit. 
5. just use these points and explain the question which user asked. 

Context: {context}
Question: {question}

Only return the helpful answer below and nothing else.
Helpful answer:
"""

def set_custom_prompt():
    """
    Prompt template for QA retrieval for each vectorstore
    """
    prompt = PromptTemplate(template=custom_prompt_template,
                            input_variables=['context', 'question'])
    return prompt

# Retrieval QA Chain
def retrieval_qa_chain(llm, prompt, db):
    qa_chain = RetrievalQA.from_chain_type(llm=llm,
                                           chain_type='stuff',
                                           retriever=db.as_retriever(search_kwargs={'k': 3}),
                                           return_source_documents=True,
                                           chain_type_kwargs={'prompt': prompt})
    return qa_chain

# Loading the model
def load_llm():
    llm = Ollama(model="llama3")
    return llm

# QA Model Function
def qa_bot():
    embeddings = OllamaEmbeddings(model="nomic-embed-text")
    db = Chroma(persist_directory=DB_CHROMA_PATH, embedding_function=embeddings)
    llm = load_llm()
    qa_prompt = set_custom_prompt()
    qa = retrieval_qa_chain(llm, qa_prompt, db)

    return qa

# Output function
def final_result(query):
    qa_result = qa_bot()
    response = qa_result({'query': query})
    return response

# Define appreciation and discouragement phrases
APPRECIATION_PHRASES = ["thanks", "thank you", "great", "good job", "well done", "appreciate it", "good"]
DISCOURAGEMENT_PHRASES = ["not helpful", "bad", "poor", "wrong", "incorrect", "disappointed"]

# Define response templates for appreciation and discouragement
APPRECIATION_RESPONSE = "You're welcome! I'm glad I could help. If you have more questions, feel free to ask."
DISCOURAGEMENT_RESPONSE = "I'm sorry to hear that you are not satisfied with my response. Could you please provide more details or ask another question so I can assist you better."

# Check if the input is an appreciation or discouragement phrase using fuzzy matching
def detect_sentiment(message):
    for phrase in APPRECIATION_PHRASES:
        if message.lower() == phrase.lower() or fuzz.ratio(message.lower(), phrase.lower()) > 90:  # 90% similarity threshold
            return 'appreciation'
    for phrase in DISCOURAGEMENT_PHRASES:
        if message.lower() == phrase.lower() or fuzz.ratio(message.lower(), phrase.lower()) > 90:  # 90% similarity threshold
            return 'discouragement'
    return None

# Define greetings and fuzzy matching for greetings
greetings = ["hi", "hello", "hey", "greetings", "good morning", "good afternoon", "good evening"]

def is_greeting(user_input):
    user_input = user_input.strip().lower()
    for greeting in greetings:
        if user_input == greeting or fuzz.ratio(user_input, greeting) > 90:  # 90% similarity threshold
            return True
    return False

# Function to check if the query is asking to generate a lesson plan
async def is_generate_lesson_plan_query(query):
    llm = load_llm()  # Load the language model
    user_query = f"""
    Please answer the following question with either "yes" or "no" only.
    query={query} 
    1. Check the query which I have provided is asking or wants us to generate a lesson plan or not. If it is asking about the lesson plan, then only return 'yes', otherwise always return 'no'.
    2. (Very Very Important - Just return only 'yes' if the user asks about generating a lesson plan else Just return only 'no' .):
    3. IF asked about images then return no. 
    """
    response = await llm.apredict(user_query)  # Ensure async call is awaited
    answer = response.strip().lower()  # Strip whitespace and convert to lowercase
    # Check if the response is exactly "yes" or "no"
    return "yes" in answer

# Global variables
table_memory = {}
lessonPlan = ""

def add_hyperlink(paragraph, url, text):
    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), paragraph.part.relate_to(url, 'rId1', is_external=True))

    # Create a new run element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add style for the hyperlink (blue and underlined)
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    # Set the text color to blue
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # Blue color
    rPr.append(color)

    # Set the underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')  # Single underline
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._element.append(hyperlink)


# Function to save response to a .docx file
async def save_response_to_docx(response, formatted_message, j_formatted_message):
    doc = Document()
    doc.add_heading('Music Blocks Lesson Plan', 0)
    
    # Add the response to the document
    doc.add_paragraph(response)
    # Add formatted_message with hyperlinks
    doc.add_paragraph("MIDI Links:")
    lines = formatted_message.splitlines()
    for line in lines:
        if line.strip():  # Avoid adding empty lines
            if '](' in line and line.endswith(')'):
                text, url = line.split('](', 1)
                text = text.replace('[', '')  # Clean up the link text
                url = url.replace(')', '')    # Clean up the URL
                p = doc.add_paragraph()
                add_hyperlink(p, url, text)
            else:
                # If the line does not match the expected format, just add it as plain text
                doc.add_paragraph(line)
    
    # Add j_formatted_message with hyperlinks
    doc.add_paragraph("Wiki Pages:")
    lines = j_formatted_message.splitlines()
    for line in lines:
        if line.strip():  # Avoid adding empty lines
            if '](' in line and line.endswith(')'):
                text, url = line.split('](', 1)
                text = text.replace('[', '')  # Clean up the link text
                url = url.replace(')', '')    # Clean up the URL
                p = doc.add_paragraph()
                add_hyperlink(p, url, text)
            else:
                # If the line does not match the expected format, just add it as plain text
                doc.add_paragraph(line)

    # Save the document
    save_path = os.path.join('lesson_plans', 'lesson_plan.docx')
    doc.save(save_path)
    return save_path 
    
# Chainlit code
@cl.on_chat_start
async def on_chat_start():
    global table_memory
    chain = qa_bot()
    msg = cl.Message(content="Starting the bot...")
    await msg.send()
    msg.content = """NOTE - Please read the Readme file so that you efficiently use our Music Blocks AI.
                     Hi, welcome to the Music Blocks lesson plan creation assistant. How may I help you?"""
    await msg.update()

    cl.user_session.set("chain", chain)
    cl.user_session.set("chat_history", [])  # Clear chat history on new chat
    cl.user_session.set("context", "")       # Clear context on new chat

def search_song(song_name, num_links):
      search_results = googlesearch.search(song_name)
      results = []
      for i, result in enumerate(search_results):
            if i >= num_links:
                break
            results.append(result)
        
      return results
  
def wiki_search(topic, num_links):
    lesson_topic = f"wiki pages of {topic}"
    search_results = googlesearch.search(lesson_topic)
    results = []
    for i, result in enumerate(search_results):
        if i>=num_links:
            break
        results.append(result)        
    return results  

# Define the system prompt for contextualization
contextualize_q_system_prompt = """
1. Given a chat history and the latest user question, formulate a standalone question.
2. Always give more priority to the latest question the user asked.
3. Ensure the reformulated question can be understood without the chat history.
4. Never answer the question; just return the updated question.
5. Do not deviate too much from the question; keep it in the same format as the user asked.
6. Take reference from the chat history if there is any relationship with the current question.
7. If the question the user asked is not related to the chat history, then just return the same latest user question. Keep it simple.
8. Don’t change the complete question; just take reference from the context.
9. If the user asks a question which is not related to Music Blocks, just return it as it is.
10. Return only the reformulated question.
11. Do not deviate too much from the question asked.
12. If the user asks maths questions return as it is.
13. If any music terms are asked try to relate with music blocks."""

# Create the prompt template for contextualization
contextualize_q_prompt = ChatPromptTemplate.from_messages(
    [
        ("system", contextualize_q_system_prompt),
        MessagesPlaceholder(variable_name="chat_history"),
        ("human", "{question}"),
    ]
)

# Combine the prompt template, LLM, and output parser into a chain
llm = load_llm()
contextualize_q_chain = contextualize_q_prompt | llm | StrOutputParser()

# Define a function to contextualize questions dynamically
def contextualize_question(chat_history, question):
    return contextualize_q_chain.invoke(
        {
            "chat_history": chat_history,
            "question": question,
        }
    )

@cl.on_message
async def main(message: cl.Message):
    global table_memory
    user_input = message.content.strip().lower()
    
    # Check for sentiment
    sentiment = detect_sentiment(user_input)
    
    if sentiment == 'appreciation':
        await cl.Message(content=APPRECIATION_RESPONSE).send()
        return
    elif sentiment == 'discouragement':
        await cl.Message(content=DISCOURAGEMENT_RESPONSE).send()
        return

    # Check for greetings
    if is_greeting(user_input):
        await cl.Message(content="Hello! How can I assist you today?").send()
        return

    # Check if user input is a request to generate or create a lesson plan
    if await is_generate_lesson_plan_query(user_input):
        global lessonPlan, topic
        lessonPlan = user_input
        user_input = f"What are the best 6 real world songs that I can use for this {user_input}? Give me the songs in a table the first column has index and second column has song(only give song name) and third column has why that song is helpful follow this order just give me a table dont give me any other information other than the table ignore all the other questions asked please"
        # Use conversation chain to form context-aware input
        # Initiate QA call to get the response
        cb = cl.AsyncLangchainCallbackHandler(
            stream_final_answer=True, answer_prefix_tokens=["FINAL", "ANSWER"]
        )
        cb.answer_reached = True
        llm = load_llm()
        result = await llm.apredict(user_input)  # Ensure async call is awaited
        answer = result
        table_memory['table'] = answer
        
        # Send the response to the user
        await cl.Message(content=answer).send()
            
        # Prompt the user to select a song index
        await cl.Message(content=f"Please give me the index of the song you want to use for generating the lesson plan.").send()
        topic = await llm.apredict(f"Extract the lesson plan topic name for this question don't use music blocks word just return the topic name. question: {lessonPlan}")
        return

    # Check if user input is a number (1-6)
    if user_input.isdigit():
        print(topic)
        selected_index = int(user_input)
        if 1 <= selected_index <= 6:
            selected_song = None
            # Extract song details from table_memory based on selected_index
            if 'table' in table_memory:
                table_content = table_memory['table']
                # | Song # | Song Name |
                lines = table_content.split('\n')
                for line in lines:
                    if line.startswith(f"| {selected_index} |"):
                        parts = line.split("|")
                        selected_song = parts[2].strip()  # Song Name
                        reason = parts[3].strip()        # Reason
                        break
                user_input = f"Question: {lessonPlan} based on this {selected_song} song? Follow the structure of lesson plans provided and try to explain the above question in detail using the song and elaborate that song using that topic using music blocks.(Important - Give step-by-step explanation of how to do it using the blocks available in music blocks use blocks which are necessary for the lessonPlan also you can use mice of Music Blocks). Explain in detail in 1000 to 1500 words the question"
                chain = cl.user_session.get("chain")
                if not chain:
                    raise ValueError("Chain is not initialized.")
                
                # Initiate QA call to get the response
                cb = cl.AsyncLangchainCallbackHandler(
                    stream_final_answer=True, answer_prefix_tokens=["FINAL", "ANSWER"]
                )
                cb.answer_reached = True
                res = await chain.acall({'query': user_input}, callbacks=[cb])  # Ensure async call is awaited
                answer = res["result"]
                
                # Send the response to the user
                await cl.Message(content=answer).send()
                midi_urls = search_song("search midi file of "+selected_song, 5)
                formatted_message = f"Here are some MIDI file links for the song '{selected_song}':\n"
                for i, url in enumerate(midi_urls, start=1):
                    formatted_message += f"[{selected_song} MIDI {i}]({url})\n"
                await cl.Message(content=formatted_message).send()
                wiki_urls = wiki_search(topic, 3)
                j_formatted_message = f"Here are some wiki pages of the topic which you can refer '{topic}':\n"
                print(f"ok {j_formatted_message}")
                for i, url in enumerate(wiki_urls, start=1):
                    j_formatted_message += f"[Wiki {i}]({url})\n"
                await cl.Message(content=j_formatted_message).send()
                # Save the response to a .docx file
                os.makedirs('lesson_plans', exist_ok=True)
                file_path = await save_response_to_docx(answer, formatted_message, j_formatted_message)
                elements = [
                cl.File(
                    name= topic,
                    path=file_path,
                    display="inline",
                ),
                ]
                await cl.Message(
                    content="Your lesson plan is ready!", elements=elements
                ).send()
                return
        else:
            if 'table' in table_memory:
                await cl.Message(content=table_memory['table']).send()
                await cl.Message(content="Please give a valid song index between 1 and 6.").send()
            else:
                await cl.Message(content="This is the music blocks chatbot.").send()
            return

    # Contextualize the question if needed
    chat_history = cl.user_session.get("chat_history", [])
    print(user_input)
    user_input = contextualize_question(chat_history, user_input)
    print(user_input)

    # If no sentiment or greeting detected, proceed with QA
    chain = cl.user_session.get("chain")
    if not chain:
        raise ValueError("Chain is not initialized.")
    cb = cl.AsyncLangchainCallbackHandler()
    cb.answer_reached = True
    res = await chain.acall({'query': user_input}, callbacks=[cb])  # Ensure async call is awaited
    answer = res["result"]

    # Update chat history
    chat_history.append(HumanMessage(content=user_input))
    chat_history.append(AIMessage(content=answer))
    cl.user_session.set("chat_history", chat_history)

    # Send the response back to the user
    await cl.Message(content=answer).send()
    
@cl.password_auth_callback
def auth_callback(username: str, password: str):
          return cl.User(
        identifier=username, metadata={"role": "user", "provider": "credentials"}
    )